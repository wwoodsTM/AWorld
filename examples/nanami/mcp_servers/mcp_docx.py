import logging
import os
import sys
import traceback
from typing import Any, Dict, List, Optional, Union

from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP
from pydantic import Field

# Import libraries for DOC/DOCX processing
# You might need to install these: pip install python-docx mammoth
try:
    import docx
except ImportError:
    logging.warning("python-docx library is not installed. DOCX processing might be limited.")
    docx = None

try:
    import mammoth
except ImportError:
    logging.warning("mammoth library is not installed. DOCX to Markdown conversion might be limited.")
    mammoth = None

# Use aworld logger (if a unified logger exists in the project)
# from aworld.logs.util import logger # Uncomment if aworld logger is available

# pylint: disable=W0707
# Initialize MCP server
mcp = FastMCP("doc-server")

# Placeholder for logger if aworld logger is not used
logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(level=logging.INFO, stream=sys.stdout)


@mcp.tool(
    description="Extracts text content from a DOC or DOCX file. Supports character limits and chunking for long documents."
)
async def extract_doc_text(
    file_path: str = Field(description="The absolute path to the DOC or DOCX file."),
    max_chars: Optional[int] = Field(
        default=None,
        description="Optional. Maximum number of characters to return. "
        "Text will be truncated if it exceeds this limit.",
    ),
    chunk_size: Optional[int] = Field(
        default=None,
        description="Optional. Size of text chunks to return (e.g., number of characters per chunk). "
        "Used when return_format is 'by_chunk'.",
    ),
    return_format: str = Field(
        default="full_text",
        description="Optional. 'full_text' to return the entire text (up to max_chars), "
        "'by_chunk' to return a list of text chunks.",
        enum=["full_text", "by_chunk"],
    ),
) -> Union[str, List[str]]:  # Return type might change based on return_format
    """
    Extracts text content from a DOC or DOCX file.

    Args:
        file_path: The absolute path to the DOC or DOCX file.
        max_chars: Optional. Maximum number of characters to return.
        chunk_size: Optional. Size of text chunks to return when return_format is 'by_chunk'.
        return_format: Optional. 'full_text' or 'by_chunk'.

    Returns:
        A string containing the full text or a list of text chunks.

    Raises:
        FileNotFoundError: If the specified file_path does not exist.
        RuntimeError: If required libraries are not installed or an error occurs.
        ValueError: If unsupported file format or invalid parameters are provided.
    """
    if docx is None:
        raise RuntimeError("python-docx library is required for DOCX processing but not found.")
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        # Implement text extraction logic here
        # Use python-docx for .docx files
        # For .doc files, you might need an additional library or conversion step
        if file_path.lower().endswith(".docx"):
            document = docx.Document(file_path)
            full_text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        elif file_path.lower().endswith(".doc"):
            # Handle .doc files (e.g., using win32com on Windows or a conversion tool)
            # This is more complex and might require platform-specific code or external tools
            raise NotImplementedError("Processing .doc files is not yet implemented.")
        else:
            raise ValueError("Unsupported file format. Only .doc and .docx are supported.")

        if max_chars is not None:
            full_text = full_text[:max_chars]

        # Implement chunking logic here if needed based on parameters
        if return_format == "by_chunk":
            if chunk_size is None or chunk_size <= 0:
                raise ValueError("chunk_size must be a positive integer when return_format is 'by_chunk'.")
            # Simple character-based chunking
            chunks = [full_text[i : i + chunk_size] for i in range(0, len(full_text), chunk_size)]
            return chunks
        elif return_format == "full_text":
            return full_text
        else:
            raise ValueError(f"Invalid return_format: {return_format}. Must be 'full_text' or 'by_chunk'.")

    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}\n{traceback.format_exc()}")
        raise RuntimeError(f"Error processing document for text extraction: {str(e)}")


@mcp.tool(description="Retrieves metadata from a DOC or DOCX file.")
async def get_doc_metadata(
    file_path: str = Field(description="The absolute path to the DOC or DOCX file."),
) -> Dict[str, Any]:
    """
    Retrieves metadata from a DOC or DOCX file.

    Args:
        file_path: The absolute path to the DOC or DOCX file.

    Returns:
        A dictionary containing metadata.

    Raises:
        FileNotFoundError: If the specified file_path does not exist.
        RuntimeError: If required libraries are not installed or an error occurs.
    """
    if docx is None:
        raise RuntimeError("python-docx library is required for DOCX processing but not found.")
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        # Implement metadata extraction logic here
        if file_path.lower().endswith(".docx"):
            document = docx.Document(file_path)
            core_properties = document.core_properties
            metadata = {
                "author": core_properties.author,
                "category": core_properties.category,
                "comments": core_properties.comments,
                "identifier": core_properties.identifier,
                "keywords": core_properties.keywords,
                "language": core_properties.language,
                "last_modified_by": core_properties.last_modified_by,
                "modified": core_properties.modified.isoformat() if core_properties.modified else None,
                "created": core_properties.created.isoformat() if core_properties.created else None,
                "revision": core_properties.revision,
                "subject": core_properties.subject,
                "title": core_properties.title,
                "version": core_properties.version,
                # Note: Page count is not reliably available in docx core properties
                # You might need to estimate or use another method if needed
            }
            return metadata
        elif file_path.lower().endswith(".doc"):
            raise NotImplementedError("Processing .doc files is not yet implemented.")
        else:
            raise ValueError("Unsupported file format. Only .doc and .docx are supported.")

    except Exception as e:
        logger.error(f"Error getting metadata from {file_path}: {e}\n{traceback.format_exc()}")
        raise RuntimeError(f"Error processing document for metadata: {str(e)}")


@mcp.tool(description="Extracts all images from a DOCX file and saves them to a specified directory.")
async def extract_doc_images(
    file_path: str = Field(description="Absolute path to the input DOCX file."),
    output_directory: str = Field(description="Directory to save the extracted images."),
) -> List[str]:
    """
    Extracts all images from a DOCX file and saves them to a specified directory.

    Args:
        file_path: Absolute path to the input DOCX file.
        output_directory: Directory to save the extracted images.

    Returns:
        A list of absolute paths to the extracted image files.

    Raises:
        RuntimeError: If required libraries are not installed.
        FileNotFoundError: If the input file does not exist.
        ValueError: If the output directory cannot be created or file format is unsupported.
        RuntimeError: If an error occurs during image extraction.
    """
    if docx is None:
        raise RuntimeError("python-docx library is required for DOCX processing but not found.")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Input file not found: {file_path}")
    if not file_path.lower().endswith(".docx"):
        raise ValueError("Unsupported file format. Only .docx is supported for image extraction.")
    if not os.path.isdir(output_directory):
        try:
            os.makedirs(output_directory, exist_ok=True)
        except Exception as e:
            raise ValueError(f"Cannot create output directory: {output_directory}, error: {str(e)}")

    extracted_image_paths = []
    try:
        document = docx.Document(file_path)
        for rel in document.part.rels.values():
            if "image" in rel.target_ref:
                image_part = rel.target_part
                image_bytes = image_part.blob
                image_ext = image_part.content_type.split("/")[-1]
                # Generate a unique filename
                image_filename = os.path.join(output_directory, f"image_{len(extracted_image_paths) + 1}.{image_ext}")
                with open(image_filename, "wb") as img_file:
                    img_file.write(image_bytes)
                extracted_image_paths.append(image_filename)

        return extracted_image_paths
    except Exception as e:
        logger.error(f"Error extracting images from {file_path}: {e}\n{traceback.format_exc()}")
        raise RuntimeError(f"Error extracting images: {str(e)}")


@mcp.tool(description="Converts a DOC or DOCX file to Markdown format.")
async def convert_doc_to_markdown(
    file_path: str = Field(description="Absolute path to the input DOC or DOCX file."),
) -> str:
    """
    Converts the content of a DOC or DOCX file into Markdown format.

    Args:
        file_path: Absolute path to the input DOC or DOCX file.

    Returns:
        A string containing the Markdown representation of the document.

    Raises:
        FileNotFoundError: If the input file does not exist.
        RuntimeError: If required libraries are not installed or an error occurs.
    """
    if mammoth is None:
        raise RuntimeError("mammoth library is required for DOCX to Markdown conversion but not found.")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Input file not found: {file_path}")
    if not (file_path.lower().endswith(".docx") or file_path.lower().endswith(".doc")):
        raise ValueError("Unsupported file format. Only .doc and .docx are supported for Markdown conversion.")

    try:
        # mammoth primarily supports .docx
        if file_path.lower().endswith(".docx"):
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                markdown_content = result.value  # The generated Markdown
                # messages = result.messages # Any messages, such as warnings during conversion
            return markdown_content
        elif file_path.lower().endswith(".doc"):
            # mammoth might have limited .doc support or require conversion first
            raise NotImplementedError(
                "Direct .doc to Markdown conversion is not reliably supported by common Python libraries."
            )
        else:
            raise ValueError("Unsupported file format. Only .doc and .docx are supported.")

    except Exception as e:
        logger.error(f"Error converting {file_path} to Markdown: {e}\n{traceback.format_exc()}")
        raise RuntimeError(f"Error converting document to Markdown: {str(e)}")


def main():
    """
    Main function to start the MCP server.
    """
    load_dotenv()  # Load environment variables
    mcp.run(transport="stdio")  # Use FastMCP's run method


# Make the module callable for uvx
def __call__():
    """
    Make the module callable for uvx.
    This function is called when the module is executed directly.
    """
    main()


# Add this for compatibility with uvx
if __name__ != "__main__":  # Ensure this is also set when imported as a module
    sys.modules[__name__].__call__ = __call__

# Run the server when the script is executed directly
if __name__ == "__main__":
    main()
